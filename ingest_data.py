import os
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
import PyPDF2
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
import time
import json
import logging
from typing import Set, List, Deque
from collections import deque
import re
import warnings
from docx import Document as DocxDocument

# Suppress warnings
warnings.filterwarnings("ignore", module="dotenv")
warnings.filterwarnings("ignore", message=".*python-dotenv.*")

logger = logging.getLogger(__name__)


class AngelOneDataIngester:
    def __init__(self):
        self.base_url = "https://www.angelone.in"
        self.support_url = "https://www.angelone.in/support"
        self.session = self._create_session()
        self.visited_urls: Set[str] = set()
        self.url_queue: Deque[str] = deque()
        self.documents: List[Document] = []
        self.max_pages = 200
        self.max_depth = 2
        self.delay_between_requests = 2
        self.ignore_robots_txt = True

        self.embeddings = self._initialize_embeddings()
        self.text_splitter = self._initialize_text_splitter()

    def _create_session(self) -> requests.Session:
        """Create and configure HTTP session"""
        session = requests.Session()
        session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
        })
        return session

    def _initialize_embeddings(self) -> HuggingFaceEmbeddings:
        """Initialize HuggingFace embeddings"""
        return HuggingFaceEmbeddings(
            encode_kwargs={'model_name': "sentence-transformers/all-MiniLM-L3-v2"}
        )

    def _initialize_text_splitter(self) -> RecursiveCharacterTextSplitter:
        """Initialize text splitter"""
        return RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )

    def _is_support_url(self, url: str) -> bool:
        """Check if URL is under the support section"""
        return url.startswith(self.support_url)

    def _get_url_depth(self, url: str) -> int:
        """Calculate the depth of a URL relative to the support base URL"""
        if not self._is_support_url(url):
            return 0

        # Remove the support base URL and count path segments
        relative_path = url[len(self.support_url):].strip('/')
        if not relative_path:
            return 0

        return len(relative_path.split('/'))

    def _should_follow_url(self, url: str) -> bool:
        """Check if we should follow this URL based on depth and other criteria"""
        if not self._is_support_url(url):
            return False

        if url in self.visited_urls:
            return False

        if self._get_url_depth(url) > self.max_depth:
            return False

        # Skip certain file types and patterns
        skip_patterns = [
            r'\.pdf$', r'\.doc$', r'\.docx$', r'\.zip$', r'\.rar$',
            r'/login', r'/register', r'/download', r'/api/',
            r'#', r'\?print=', r'/print/', r'javascript:'
        ]

        for pattern in skip_patterns:
            if re.search(pattern, url, re.IGNORECASE):
                return False

        return True

    def _extract_all_links(self, soup: BeautifulSoup, current_url: str) -> List[str]:
        """Extract all valid support links from the current page"""
        links = []

        # Find all anchor tags with href attributes
        for link in soup.find_all('a', href=True):
            href = link['href']

            # Skip empty hrefs and javascript
            if not href or href.startswith('javascript:') or href.startswith('mailto:'):
                continue

            # Convert relative URLs to absolute URLs
            full_url = urljoin(current_url, href)

            # Clean fragment identifiers and query parameters for consistency
            clean_url = full_url.split('#')[0].split('?')[0]

            # Check if we should follow this URL
            if self._should_follow_url(clean_url):
                links.append(clean_url)

        return list(set(links))  # Remove duplicates

    def _remove_unwanted_elements(self, soup: BeautifulSoup) -> None:
        """Remove unwanted HTML elements from soup"""
        unwanted_tags = [
            "script", "style", "nav", "footer", "header", "aside",
            "button", "input", "form", "advertisement", "ad"
        ]
        for element in soup(unwanted_tags):
            element.decompose()

    def _find_content_areas(self, soup: BeautifulSoup) -> List:
        """Find main content areas in the soup"""
        content_selectors = [
            'main', 'article', '[role="main"]',
            '.content', '.main-content', '.support-content',
            '.faq-content', '.help-content', '#main-content',
            '.container', '.wrapper'
        ]

        for selector in content_selectors:
            elements = soup.select(selector)
            if elements:
                return elements

        return [soup.find('body')] if soup.find('body') else []

    def _extract_text_from_areas(self, content_areas: List) -> str:
        """Extract text content from content areas"""
        text_content = ""

        for area in content_areas:
            if area:
                for element in area.find_all(['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li']):
                    text = element.get_text(strip=True)
                    if text and len(text) > 10:
                        text_content += text + "\n"

        return text_content

    def _create_unique_filename(self, base_filename: str) -> str:
        """Create a unique filename if file already exists"""
        filename = base_filename
        counter = 1

        while os.path.exists(filename):
            base, ext = os.path.splitext(base_filename)
            filename = f"{base}_{counter}{ext}"
            counter += 1

        return filename

    def _save_page_content(self, url: str, title: str, content: str) -> None:
        """Save page content to file"""
        parsed = urlparse(url)
        safe_path = parsed.path.replace('/', '_').replace('?', '_').replace('&', '_').replace('#', '_')
        if not safe_path or safe_path == '_':
            safe_path = "main_page"

        base_filename = f"data/scraped_pages/{safe_path}.txt"
        filename = self._create_unique_filename(base_filename)

        with open(filename, 'w', encoding='utf-8') as f:
            f.write(f"Title: {title}\nURL: {url}\nDepth: {self._get_url_depth(url)}\n\n{content}")

    def _is_valid_content(self, text: str) -> bool:
        """Check if the extracted text is valid content"""
        if len(text.strip()) < 50:
            return False

        words = text.lower().split()
        if not words:
            return False

        common_nav_words = ['menu', 'navigation', 'footer', 'header', 'login', 'register', 'sign up', 'sign in']
        nav_word_count = sum(1 for word in words if word in common_nav_words)

        return nav_word_count / len(words) < 0.1

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)

        # Remove common boilerplate phrases
        boilerplate_patterns = [
            r'©.*?\d{4}.*?Angel One',
            r'All rights reserved',
            r'Terms and Conditions',
            r'Privacy Policy',
            r'Cookie Policy',
        ]

        for pattern in boilerplate_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)

        return text.strip()

    def _crawl_page(self, url: str):
        """Crawl a single page and extract content, then find new URLs to crawl"""
        if url in self.visited_urls or len(self.visited_urls) >= self.max_pages:
            return

        # Strict check: only process URLs under /support
        if not self._is_support_url(url):
            return

        try:
            depth = self._get_url_depth(url)
            print(f"Scraping (depth {depth}): {url}")

            response = self.session.get(url, timeout=15)
            response.raise_for_status()

            self.visited_urls.add(url)
            soup = BeautifulSoup(response.content, 'html.parser')

            # Extract title
            title = soup.find('title')
            title_text = title.get_text().strip() if title else "No Title"

            # Extract new URLs to crawl before removing unwanted elements
            new_urls = self._extract_all_links(soup, url)

            # Add new URLs to queue
            for new_url in new_urls:
                if new_url not in self.visited_urls and len(self.visited_urls) + len(self.url_queue) < self.max_pages:
                    self.url_queue.append(new_url)

            # Remove unwanted elements for content extraction
            self._remove_unwanted_elements(soup)

            # Find and extract content
            content_areas = self._find_content_areas(soup)
            text_content = self._extract_text_from_areas(content_areas)

            # Clean and validate content
            cleaned_text = self._clean_text(text_content)

            if self._is_valid_content(cleaned_text):
                # Save content to file
                self._save_page_content(url, title_text, cleaned_text)

                # Create document
                doc = Document(
                    page_content=cleaned_text,
                    metadata={
                        "source": url,
                        "title": title_text,
                        "type": "web_page",
                        "depth": depth,
                        "content_length": len(cleaned_text)
                    }
                )
                self.documents.append(doc)
                print(f"Successfully extracted content from {url} (found {len(new_urls)} new URLs)")
            else:
                print(f"Invalid content for {url}, skipping...")

        except requests.exceptions.RequestException as e:
            print(f"Request error for {url}: {str(e)}")
        except Exception as e:
            print(f"Error scraping {url}: {str(e)}")

    def scrape_support_pages(self):
        """Scrape support pages recursively using breadth-first search"""
        print("Starting to scrape Angel One support pages (multi-layer)...")
        print(f"Max depth: {self.max_depth}, Max pages: {self.max_pages}")

        if self.ignore_robots_txt:
            print("WARNING: Ignoring robots.txt restrictions!")

        os.makedirs("data/scraped_pages", exist_ok=True)

        # Start with the main support page
        self.url_queue.append(self.support_url)

        # Process URLs in breadth-first manner
        while self.url_queue and len(self.visited_urls) < self.max_pages:
            current_url = self.url_queue.popleft()

            if current_url not in self.visited_urls:
                # Add delay between requests
                if self.visited_urls:  # Don't delay for the first request
                    time.sleep(self.delay_between_requests)

                self._crawl_page(current_url)

        print(f"Scraping completed!")
        print(f"Total pages scraped: {len(self.visited_urls)}")
        print(f"URLs remaining in queue: {len(self.url_queue)}")

        # Print depth distribution
        depth_counts = {}
        for doc in self.documents:
            depth = doc.metadata.get('depth', 0)
            depth_counts[depth] = depth_counts.get(depth, 0) + 1

        print("Pages scraped by depth:")
        for depth in sorted(depth_counts.keys()):
            print(f"  Depth {depth}: {depth_counts[depth]} pages")

    def _extract_pdf_text(self, pdf_path: str) -> tuple[str, int]:
        """Extract text from a single PDF file"""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text_content = ""

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content += page_text + "\n"
                except Exception as e:
                    print(f"Warning: Could not extract text from page {page_num + 1}: {e}")

            return text_content, len(pdf_reader.pages)

    def _process_single_pdf(self, filename: str, pdf_dir: str) -> None:
        """Process a single PDF file"""
        pdf_path = os.path.join(pdf_dir, filename)

        try:
            text_content, page_count = self._extract_pdf_text(pdf_path)

            if len(text_content.strip()) > 100:
                cleaned_text = self._clean_text(text_content)

                doc = Document(
                    page_content=cleaned_text,
                    metadata={
                        "source": filename,
                        "type": "pdf_document",
                        "pages": page_count,
                        "content_length": len(cleaned_text)
                    }
                )
                self.documents.append(doc)
                print(f"Successfully processed {filename} ({page_count} pages)")

        except Exception as e:
            print(f"Error processing PDF {filename}: {str(e)}")

    def process_pdfs(self):
        """Process PDF files in the data/pdfs directory"""
        print("Processing PDF files...")

        pdf_dir = "data/pdfs"
        pdf_files = [f for f in os.listdir(pdf_dir) if f.lower().endswith('.pdf')]
        if not pdf_files:
            return
        for filename in pdf_files:
            self._process_single_pdf(filename, pdf_dir)

    def _extract_docx_text(self, docx_path: str) -> tuple[str, int, int]:
        """Extract text from a single DOCX file"""
        docx_doc = DocxDocument(docx_path)
        text_content = ""

        # Extract text from paragraphs
        for paragraph in docx_doc.paragraphs:
            if paragraph.text.strip():
                text_content += paragraph.text + "\n"

        # Extract text from tables
        for table in docx_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content += cell.text + " "
                text_content += "\n"

        # Extract text from headers and footers
        for section in docx_doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        text_content += paragraph.text + "\n"

            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        text_content += paragraph.text + "\n"

        paragraph_count = len([p for p in docx_doc.paragraphs if p.text.strip()])
        table_count = len(docx_doc.tables)

        return text_content, paragraph_count, table_count

    def _process_single_docx(self, filename: str, docx_dir: str) -> None:
        """Process a single DOCX file"""
        docx_path = os.path.join(docx_dir, filename)

        try:
            text_content, paragraph_count, table_count = self._extract_docx_text(docx_path)

            if len(text_content.strip()) > 100:
                cleaned_text = self._clean_text(text_content)

                doc = Document(
                    page_content=cleaned_text,
                    metadata={
                        "source": filename,
                        "type": "docx_document",
                        "paragraphs": paragraph_count,
                        "tables": table_count,
                        "content_length": len(cleaned_text),
                        "found_in_directory": docx_dir
                    }
                )
                self.documents.append(doc)
                print(f"Successfully processed {filename} ({paragraph_count} paragraphs, {table_count} tables)")

        except Exception as e:
            print(f"Error processing DOCX {filename}: {str(e)}")

    def process_docx_files(self):
        """Process DOCX files in the data/docx directory"""
        print("Processing DOCX files...")

        docx_dir = "data/docx"
        docx_files = [f for f in os.listdir(docx_dir) if f.lower().endswith('.docx')]
        if not docx_files:
            return
        for filename in docx_files:
            self._process_single_docx(filename, docx_dir)

    def process_documents(self):
        """Process both PDF and DOCX files in the data directory"""
        print("Processing document files...")
        self.process_pdfs()
        self.process_docx_files()

    def _create_document_chunks(self) -> List[Document]:
        """Create chunks from all documents"""
        all_chunks = []

        for doc in self.documents:
            try:
                chunks = self.text_splitter.split_text(doc.page_content)
                for i, chunk in enumerate(chunks):
                    if len(chunk.strip()) > 50:  # Only include substantial chunks
                        chunk_metadata = doc.metadata.copy()
                        chunk_metadata['chunk_id'] = i
                        chunk_doc = Document(
                            page_content=chunk.strip(),
                            metadata=chunk_metadata
                        )
                        all_chunks.append(chunk_doc)
            except Exception as e:
                print(f"Warning: Could not split document {doc.metadata.get('source', 'unknown')}: {e}")

        return all_chunks

    def _save_vector_store_metadata(self, persist_directory: str, total_chunks: int) -> None:
        """Save metadata about the vector store"""
        metadata = {
            "total_documents": len(self.documents),
            "total_chunks": total_chunks,
            "sources": list(set([doc.metadata.get("source", "unknown") for doc in self.documents])),
            "document_types": list(set([doc.metadata.get("type", "unknown") for doc in self.documents])),
            "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "scraping_summary": {
                "urls_visited": len(self.visited_urls),
                "web_pages_processed": len([d for d in self.documents if d.metadata.get('type') == 'web_page']),
                "pdfs_processed": len([d for d in self.documents if d.metadata.get('type') == 'pdf_document']),
                "docx_processed": len([d for d in self.documents if d.metadata.get('type') == 'docx_document']),
                "max_depth": self.max_depth,
                "robots_txt_ignored": self.ignore_robots_txt,
                "restriction": "Only URLs under https://www.angelone.in/support"
            }
        }

        with open(f"{persist_directory}/metadata.json", "w") as f:
            json.dump(metadata, f, indent=2)

    def create_vector_store(self):
        """Create and save Chroma vector store"""
        print("Creating vector store...")

        if not self.documents:
            print("No documents to process!")
            return None

        # Create chunks
        all_chunks = self._create_document_chunks()
        print(f"Created {len(all_chunks)} chunks from {len(self.documents)} documents")

        if not all_chunks:
            print("No chunks created. Cannot create vector store.")
            return None

        # Create Chroma vector store
        persist_directory = "chroma_db"

        try:
            print("Building vector store (this may take a few minutes)...")

            vector_store = Chroma.from_documents(
                documents=all_chunks,
                embedding=self.embeddings,
                persist_directory=persist_directory
            )

            os.makedirs(persist_directory, exist_ok=True)
            self._save_vector_store_metadata(persist_directory, len(all_chunks))

            print("Vector store created and saved successfully!")
            return vector_store

        except Exception as e:
            print(f"Error creating vector store: {e}")
            return None

    def _print_ingestion_summary(self) -> None:
        """Print summary of the ingestion process"""
        print("\n" + "=" * 60)
        print("INGESTION SUMMARY")
        print("=" * 60)
        print(f"Total documents processed: {len(self.documents)}")
        print(f"Web pages scraped: {len([d for d in self.documents if d.metadata.get('type') == 'web_page'])}")
        print(f"PDF files processed: {len([d for d in self.documents if d.metadata.get('type') == 'pdf_document'])}")
        print(f"DOCX files processed: {len([d for d in self.documents if d.metadata.get('type') == 'docx_document'])}")
        print(f"URLs visited: {len(self.visited_urls)}")
        print(
            f"Max depth reached: {max([d.metadata.get('depth', 0) for d in self.documents if d.metadata.get('type') == 'web_page'], default=0)}")
        print("=" * 60)

    def run_ingestion(self):
        """Run the complete data ingestion pipeline"""
        print("Starting RAG data ingestion...")
        print("=" * 60)

        # Step 1: Scrape support pages recursively
        self.scrape_support_pages()

        # Step 2: Process documents (PDFs and DOCX)
        self.process_documents()

        if not self.documents:
            print("No documents were processed. Please check your data directories.")
            return None

        # Step 3: Create vector store
        vector_store = self.create_vector_store()

        # Print summary
        self._print_ingestion_summary()

        if vector_store:
            print(f"Vector store saved to: chroma_db/")
            print("Ready to start the RAG system!")
        else:
            print("Vector store creation failed")

        return vector_store


def setup_directories():
    """Create necessary data directories"""
    directories = ["data/pdfs", "data/docx", "data/scraped_pages"]
    for directory in directories:
        os.makedirs(directory, exist_ok=True)


if __name__ == "__main__":
    print("Setting up environment...")

    # Create data directories
    setup_directories()

    # Run ingestion
    try:
        ingester = AngelOneDataIngester()
        vector_store = ingester.run_ingestion()

        if not vector_store:
            print("\nIngestion failed. Please check the error messages above.")

    except KeyboardInterrupt:
        print("\nIngestion stopped by user")
    except Exception as e:
        print(f"\nUnexpected error during ingestion: {e}")